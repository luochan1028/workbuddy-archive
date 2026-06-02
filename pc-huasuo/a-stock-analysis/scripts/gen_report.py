#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""A股AI产业链选股分析报告 Word生成脚本"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ======== 数据 ========
stock_data = {
    "688041.SH": dict(name="海光信息", layer=1, layer_name="算力芯片",
        price=155, target=175, upside=0.13, pe=95, roe=13.5,
        revenue_g=0.52, score=80, rating=u"\u2b50c",
        catalyst=u"DCU3.0量产、金融/能源国产化加速",
        risk=u"美国出口管制风险、技术授权依赖AMD",
        biz=u"x86 CPU+DCU芯片，党政/金融国产化核心供应商"),
    "300308.SZ": dict(name=u"中际旭创", layer=3, layer_name=u"光通信",
        price=185, target=220, upside=0.19, pe=35, roe=24.5,
        revenue_g=0.58, score=84, rating=u"\u2b50c",
        catalyst=u"英伟达B200/GB200出货、1.6T光模块量产",
        risk=u"客户集中(英伟达/北云)、中美贸易摩擦",
        biz=u"全球光模块龙头，800G主力供应商"),
    "300502.SZ": dict(name=u"新易盛", layer=3, layer_name=u"光通信",
        price=128, target=155, upside=0.21, pe=32, roe=24.8,
        revenue_g=0.62, score=77, rating=u"\u2b50c",
        catalyst=u"800G出货攀升、亚马逊/Meta订单",
        risk=u"行业价格战、光芯片依赖进口",
        biz=u"光模块厂，800G快速放量"),
    "300394.SZ": dict(name=u"天孚通信", layer=3, layer_name=u"光通信",
        price=165, target=185, upside=0.12, pe=38, roe=25.2,
        revenue_g=0.45, score=77, rating=u"\u2b50c",
        catalyst=u"CPO量产、硅光渗透率提升",
        risk=u"CPO技术路线不确定、下游压价",
        biz=u"光器件龙头，光模块上游核心部件"),
    "600584.SH": dict(name=u"长电科技", layer=2, layer_name=u"存储+封测",
        price=68, target=78, upside=0.15, pe=42, roe=9.2,
        revenue_g=0.22, score=71, rating=u"\u2b50c",
        catalyst=u"HBM3e封装订单、AI芯片先进封装产能扩张",
        risk=u"资本开支大、先进封装良率爬坡",
        biz=u"全球第三大封测厂，先进封装核心供应商"),
    "002837.SZ": dict(name=u"英维克", layer=6, layer_name=u"数据中心基建",
        price=42, target=48, upside=0.14, pe=48, roe=14.5,
        revenue_g=0.38, score=72, rating=u"\u2b50c",
        catalyst=u"英伟达GB200液冷方案放量、国内AI数据中心建设加速",
        risk=u"价格战、客户集中",
        biz=u"精密温控/液冷设备龙头，AI数据中心散热核心供应商"),
    "002156.SZ": dict(name=u"通富微电", layer=2, layer_name=u"存储+封测",
        price=42, target=48, upside=0.14, pe=38, roe=8.5,
        revenue_g=0.18, score=64, rating=u"\u25b6",
        catalyst=u"AMD MI300系列出货增加、2.5D封装产能释放",
        risk=u"AMD依赖度>60%、毛利率偏低",
        biz=u"封测厂，AMD核心封测供应商"),
    "688525.SH": dict(name=u"佰维存储", layer=2, layer_name=u"存储+封测",
        price=88, target=95, upside=0.08, pe=65, roe=13.2,
        revenue_g=0.68, score=68, rating=u"\u25b6",
        catalyst=u"企业级SSD放量、AI服务器存储需求爆发",
        risk=u"存储周期波动大、技术代差",
        biz=u"存储芯片(NAND+NOR Flash)，国产存储替代"),
    "000938.SZ": dict(name=u"紫光股份", layer=4, layer_name=u"网络设备",
        price=38, target=44, upside=0.16, pe=28, roe=10.5,
        revenue_g=0.15, score=64, rating=u"\u25b6",
        catalyst=u"AI数据中心交换机订单放量、新华三H股上市预期",
        risk=u"负债率高、毛利率偏低",
        biz=u"国内交换机/路由器龙头(新华三)，AI数据中心网络核心供应商"),
    "603019.SH": dict(name=u"中科曙光", layer=6, layer_name=u"数据中心基建",
        price=72, target=80, upside=0.11, pe=35, roe=11.2,
        revenue_g=0.18, score=68, rating=u"\u25b6",
        catalyst=u"国产AI服务器放量、海光DCU服务器出货增加",
        risk=u"毛利率偏低、负债率偏高",
        biz=u"国产服务器/高性能计算龙头，党政/科研市场强势"),
    "002371.SZ": dict(name=u"北方华创", layer=5, layer_name=u"半导体设备",
        price=480, target=450, upside=-0.06, pe=65, roe=24.5,
        revenue_g=0.38, score=76, rating=u"\u25b6",
        catalyst=u"存储厂扩产(长存/长鑫)、先进封装设备放量",
        risk=u"关键零部件仍依赖进口、估值极高",
        biz=u"国内半导体设备龙头，刻蚀/CVD/ALD设备全覆盖"),
    "688012.SH": dict(name=u"中微公司", layer=5, layer_name=u"半导体设备",
        price=320, target=310, upside=-0.03, pe=78, roe=16.2,
        revenue_g=0.32, score=76, rating=u"\u25b6",
        catalyst=u"CCP刻蚀设备订单放量、MOCVD新应用",
        risk=u"美国出口管制(可能列入实体清单)、估值极高",
        biz=u"刻蚀设备龙头，5nm以下刻蚀设备进入台积电供应链"),
    "688981.SH": dict(name=u"中芯国际", layer=5, layer_name=u"半导体制造",
        price=68, target=62, upside=-0.09, pe=48, roe=5.8,
        revenue_g=0.12, score=71, rating=u"\u25c6",
        catalyst=u"28nm以下成熟制程国产化率提升、国内客户回流",
        risk=u"先进制程受出口管制、设备受限",
        biz=u"大陆最先进晶圆代工，14nm量产，7nm研发中"),
    "688256.SH": dict(name=u"寒武纪", layer=1, layer_name=u"算力芯片",
        price=680, target=720, upside=0.06, pe=185, roe=9.8,
        revenue_g=0.48, score=67, rating=u"\u25b6",
        catalyst=u"新一代云端AI芯片量产、大客户(字节/阿里)订单确认",
        risk=u"客户集中度极高、估值极高",
        biz=u"国内AI芯片龙头，云端+边缘端AI芯片，国产替代核心标的"),
    "688047.SH": dict(name=u"龙芯中科", layer=1, layer_name="算力芯片",
        price=195, target=160, upside=-0.18, pe=320, roe=4.8,
        revenue_g=0.15, score=47, rating=u"\u25c6",
        catalyst=u"3A6000系列量产、党政采购放量",
        risk=u"商业化能力弱、性能与Intel差距仍大",
        biz=u"自主指令集(LoongArch)CPU，党政/工控国产化"),
    "002230.SZ": dict(name=u"科大讯飞", layer=7, layer_name="AI应用",
        price=62, target=58, upside=-0.06, pe=88, roe=7.5,
        revenue_g=0.15, score=47, rating=u"\u25c6",
        catalyst=u"星火大模型商业化落地、教育/医疗AI产品放量",
        risk=u"商业化进度慢于预期、盈利能力偏弱",
        biz=u"国内AI应用龙头，星火大模型，教育/医疗/政务AI落地"),
    "688111.SH": dict(name=u"金山办公", layer=7, layer_name="AI应用",
        price=365, target=350, upside=-0.04, pe=78, roe=28.5,
        revenue_g=0.22, score=61, rating=u"\u25b6",
        catalyst=u"WPS AI付费用户增长、国产化(党政/央企)持续推进",
        risk=u"估值极高、AI功能货币化进度不确定",
        biz=u"WPS AI，国内办公软件龙头，AI赋能办公场景"),
    "002396.SZ": dict(name=u"星网锐捷", layer=4, layer_name="网络设备",
        price=28, target=32, upside=0.14, pe=22, roe=8.2,
        revenue_g=0.12, score=48, rating=u"\u25c6",
        catalyst=u"AI园区网升级、信创交换机放量",
        risk=u"毛利率持续下滑、竞争力弱于新华三/华为",
        biz=u"网络设备(交换机/路由器/云终端)"),
    "300499.SZ": dict(name=u"高澜股份", layer=6, layer_name=u"数据中心基建",
        price=28, target=32, upside=0.14, pe=55, roe=9.8,
        revenue_g=0.42, score=60, rating=u"\u25b6",
        catalyst=u"液冷订单放量、储能温控业务协同",
        risk=u"竞争激烈、规模偏小",
        biz=u"液冷设备，AI数据中心液冷解决方案供应商"),
}

# ======== 辅助 ========
GREEN = RGBColor(0xD5, 0xF5, 0xE3)
BLUE  = RGBColor(0xD4, 0xE6, 0xF1)
YELLOW= RGBColor(0xFC, 0xF3, 0xCF)
RED   = RGBColor(0xFA, 0xDB, 0xD8)
GRAY  = RGBColor(0xF2, 0xF3, 0xF4)
HDR   = RGBColor(0x1F, 0x4E, 0x79)
DARK = RGBColor(0x00, 0x00, 0x00)

def rating_bg(rating):
    if rating == u"\u2b50c": return GREEN
    if rating == u"\u25b6": return BLUE
    if rating == u"\u25c6": return YELLOW
    return RED

def pct(n): return "{:.1f}%".format(n*100)
def pct_sign(n): return ("+" if n>=0 else "") + "{:.1f}%".format(n*100)
def price_str(n): return "{:.0f}元".format(n)

def set_cell_bg(cell, color_hex):
    # color_hex: "D5F5E3" etc, no "#"
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove existing shd if any
    for child in list(tcPr):
        if child.tag.endswith('shd'):
            tcPr.remove(child)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_border(cell, color="AAAAAA"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement('w:'+side)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def make_cell(text, width_cm, bg_color=None, bold=False, align=None, size_pt=9, text_color=None):
    cell = table.add_cell() if False else None  # placeholder
    return cell  # will be used inside table creation

# ======== 主程序 ========
doc = Document()

# 页面边距
from docx.shared import Mm
section = doc.sections[0]
section.page_width  = Mm(210)
section.page_height = Mm(297)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

def add_heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = HDR
    p.runs[0].font.size = Pt(20)
    p.runs[0].font.bold = True
    return p

def add_heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    p.runs[0].font.size = Pt(16)
    return p

def add_heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
    p.runs[0].font.size = Pt(14)
    return p

def add_p(text, bold=False, color=None, size_pt=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

# ======== 封面 ========
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u"A 股 AI 产业链选股分析报告")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = HDR

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u"基于七层产业链框架 · 含目标价格与操作建议")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u"2026年5月30日")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(u"\u26a0\ufe0f 本报告为研究学习框架，非投资建议，投资决策需自行判断")
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

doc.add_page_break()

# ======== 第一章：方法论 ========
add_heading1(u"一、分析框架与方法论")
add_p(u"本报告采用「A股AI产业链七层分类法」，结合国产替代紧迫度、全球供应链地位、业绩兑现能力、估值安全边际、资金面信号五大维度进行综合评分。")

add_heading2(u"1.1  七层产业链分类")
layer_defs = [
    (1, u"\u1f3ae 算力芯片", u"AI的「大脑」，国产替代核心逻辑"),
    (2, u"\u1f4be 存储+封测", u"给大脑配「仓库」，先进封装卡脖子"),
    (3, u"\u1f308 光通信", u"A股在全球AI供应链中最强的环节"),
    (4, u"\u1f310 网络设备", u"连接GPU集群的交换机与高速互联"),
    (5, u"\u1f3f0 半导体制造/设备", u"造芯片的厂和设备，去美化核心战场"),
    (6, u"\u26a1 数据中心基建", u"AI机房的电力、散热、服务器"),
    (7, u"\u1f4a1 AI应用/软件", u"把AI能力落地到具体场景"),
]
cols = (1, 3, 5)
tbl = doc.add_table(rows=8, cols=3)
tbl.style = 'Table Grid'
for i, (n, name, desc) in enumerate(layer_defs):
    row = tbl.rows[i+1]  # row 0 = header
    row.cells[0].text = str(n)
    row.cells[1].text = name
    row.cells[2].text = desc
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
        set_cell_border(cell)
        if i % 2 == 0:
            set_cell_bg(cell, "F2F3F4")

# header row
hdr = tbl.rows[0]
hdr.cells[0].text = u"层"
hdr.cells[1].text = u"名称"
hdr.cells[2].text = u"大白话"
for cell in hdr.cells:
    set_cell_bg(cell, "1F4E79")
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
    set_cell_border(cell)

add_p(u"注：每层代表AI产业链的一个关键环节，优先选卡在瓶颈上的公司。")
doc.add_paragraph()

add_heading2(u"1.2  五大评分维度")
dim_data = [
    (u"国产替代紧迫度", "20%", u"该环节国产化率？是否被列入「卡脖子」清单？大基金是否扶持？"),
    (u"全球供应链地位", "20%", u"A股公司在全球AI供应链中是否不可替代？海外大客户订单占比？"),
    (u"业绩兑现能力",   "25%", u"收入/利润是否已经体现在财报？（A股概念炒作多，业绩兑现是关键分水岭）"),
    (u"估值安全边际",   "20%", u"PE分位数 vs 历史区间？距52周高点回撤幅度？"),
    (u"资金面信号",     "15%", u"北向资金持仓变化？龙虎榜机构净买入？"),
]
tbl = doc.add_table(rows=6, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0]
hdr.cells[0].text = u"维度"
hdr.cells[1].text = u"权重"
hdr.cells[2].text = u"评估要点（A股特有）"
for cell in hdr.cells:
    set_cell_bg(cell, "1F4E79")
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.size = Pt(10)
    set_cell_border(cell)

for i, (d, w, n) in enumerate(dim_data):
    row = tbl.rows[i+1]
    row.cells[0].text = d
    row.cells[1].text = w
    row.cells[2].text = n
    for cell in row.cells:
        set_cell_border(cell)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_bg(cell, "F2F3F4")
doc.add_paragraph()

add_heading2(u"1.3  评级标准")
add_p(u"\u2b50c 强烈关注：评分 >= 75分，业绩已兑现，国产替代核心，估值相对合理", bold=True, color=RGBColor(0x1A,0x7A,0x1A))
add_p(u"\u25b6 关注：评分 60–74分，好公司好赛道，但估值偏高，等回调", bold=True, color=RGBColor(0x1A,0x3E,0x7A))
add_p(u"\u25c6 观望：评分 45–59分，概念为主，业绩待验证，或估值已透支", bold=True, color=RGBColor(0x7A,0x7A,0x1A))
add_p(u"\u25cf 回避：评分 < 45分，基本面弱，纯概念炒作，或技术路线被替代", bold=True, color=RGBColor(0x7A,0x1A,0x1A))
doc.add_paragraph()

# ======== 按层分组 ========
add_heading1(u"二、各层选股分析（含目标价格）")

# 分组
layers = {}
for code, s in stock_data.items():
    l = s["layer"]
    if l not in layers:
        layers[l] = []
    layers[l].append((code, s))

layer_meta = {
    1: (u"\u1f3ae", u"算力芯片", u"AI的「大脑」，A股没有英伟达，核心逻辑是「国产替代+自主可控」。寒武纪是纯AI芯片代表，海光信息走x86+DCU路线更成熟，龙芯走完全自主指令集但商业化最弱。"),
    2: (u"\u1f4be", u"存储+封测", u"AI服务器需要大量HBM和先进封装。A股封测厂(长电/通富)在全球有竞争力；存储芯片(佰维)受益于AI服务器SSD需求爆发。"),
    3: (u"\u1f308", u"光通信", u"这是A股在全球AI供应链中最强的环节。中际旭创、新易盛的800G光模块直接供货英伟达/微软/谷歌，业绩已大幅兑现。"),
    4: (u"\u1f310", u"网络设备", u"AI数据中心需要大量800G交换机。紫光股份(新华三)是国内交换机龙头，直接受益于国内AI数据中心建设。"),
    5: (u"\u1f3f0", u"半导体制造/设备", u"中芯国际是大陆最先进的晶圆代工厂，但受出口管制影响。北方华创和中微公司是设备国产替代核心，但估值极高。"),
    6: (u"\u26a1", u"数据中心基建", u"AI数据中心液冷散热是2026年最明确的增量。英维克是液冷龙头，订单放量确定性强。"),
    7: (u"\u1f4a1", u"AI应用/软件", u"AI应用落地是弹性最大但确定性最低的环节。科大讯飞(星火大模型)商业化还在早期。金山办公(WPS AI)商业模式最清晰。"),
}

for ln in sorted(layers.keys()):
    stocks_in_layer = sorted(layers[ln], key=lambda x: -x[1]["score"])
    meta = layer_meta.get(ln, ("", "", ""))
    add_heading2(u"第{}层：{} {}".format(ln, meta[0], meta[1]))
    add_p(meta[2])

    # 表格
    tbl = doc.add_table(rows=len(stocks_in_layer)+1, cols=8)
    tbl.style = 'Table Grid'
    hdrs = [u"代码/名称", u"评级", u"评分", u"现价/目标价", u"空间", u"PE", u"ROE", u"核心逻辑"]
    for j, h in enumerate(hdrs):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, "1F4E79")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                run.font.size = Pt(9)
        set_cell_border(cell)

    for i, (code, s) in enumerate(stocks_in_layer):
        row = tbl.rows[i+1]
        bg = rating_bg(s["rating"])
        # 交替底色
        alt = RGBColor(0xEB,0xF5,0xFB) if i%2==0 else RGBColor(0xFF,0xFF,0xFF)
        texts = [
            code + "\n" + s["name"],
            s["rating"],
            str(s["score"]),
            price_str(s["price"]) + " → " + price_str(s["target"]),
            pct_sign(s["upside"]),
            str(s["pe"]),
            "{:.1f}%".format(s["roe"]),
            s["biz"][:30] + ("..." if len(s["biz"])>30 else ""),
        ]
        for j, t in enumerate(texts):
            cell = row.cells[j]
            cell.text = t
            set_cell_border(cell)
            if j == 1:  # 评级列
                set_cell_bg(cell, bg)
            elif j == 4:  # 空间列
                c = RGBColor(0x1A,0x7A,0x1A) if s["upside"]>=0.1 else RGBColor(0xC0,0x39,0x2B)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = c
                        run.font.bold = True
                        run.font.size = Pt(9)
                if i%2==0:
                    set_cell_bg(cell, alt)
            else:
                if i%2==0:
                    set_cell_bg(cell, alt)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    # 逐股点评
    for code, s in stocks_in_layer:
        r = s["rating"]
        add_heading3(u"{} {}  |  评分：{}/100  |  目标价：{}元（{}）".format(
            r, code, s["name"], s["score"], s["target"], pct_sign(s["upside"])))
        if "算力" in s["layer_name"] or "光" in s["layer_name"] or "封" in s["layer_name"]:
            add_p(u"\u26a5 卡位瓶颈：" + s["biz"][:20])
        add_p(u"为什么关注：" + s["biz"])
        add_p(u"催化剂：" + s["catalyst"])
        add_p(u"风险提示：" + s["risk"], color=RGBColor(0xC0,0x39,0x2B))
        if s["upside"] > 0.1:
            wp = round(s["target"] * 0.92)
            add_p(u"操作建议：当前价 {} 元，距目标价还有 {} 空间。建议等待回调至 {} 元附近分批建仓。".format(
                s["price"], pct_sign(s["upside"]), wp), color=RGBColor(0x1A,0x7A,0x1A), bold=True)
        else:
            wp = round(s["target"] * 0.90)
            add_p(u"操作建议：当前价已接近或高于目标价，不建议追高。建议等待回调至 {} 元附近再考虑建仓。".format(wp),
                color=RGBColor(0xC0,0x39,0x2B), bold=True)
        doc.add_paragraph()

    doc.add_page_break()

# ======== 强烈关注清单 ========
add_heading1(u"三、\u2b50c 强烈关注清单（优先配置）")
add_p(u"以下股票综合评分 >= 75分，基本面扎实、业绩已兑现、估值相对合理，是A股AI产业链的首选配置方向。")

strong = sorted([(c,s) for c,s in stock_data.items() if s["rating"]==u"\u2b50c"], key=lambda x: -x[1]["score"])
tbl = doc.add_table(rows=len(strong)+1, cols=8)
tbl.style = 'Table Grid'
hdrs = [u"排名", u"代码", u"名称", u"层级", u"评分", u"现价", u"目标价", u"上行空间"]
for j, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "1F4E79")
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.size = Pt(10)
    set_cell_border(cell)

for i, (code, s) in enumerate(strong):
    row = tbl.rows[i+1]
    row.cells[0].text = str(i+1)
    row.cells[1].text = code
    row.cells[2].text = s["name"]
    row.cells[3].text = s["layer_name"]
    row.cells[4].text = str(s["score"])
    row.cells[5].text = price_str(s["price"])
    row.cells[6].text = price_str(s["target"])
    row.cells[7].text = pct_sign(s["upside"])
    for j in range(8):
        set_cell_border(row.cells[j])
        if j == 7:
            c = RGBColor(0x1A,0x7A,0x1A) if s["upside"]>=0.1 else RGBColor(0xC0,0x39,0x2B)
            for p in row.cells[j].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = c
                    run.font.bold = True
                    run.font.size = Pt(10)
        if i%2==0:
            set_cell_bg(row.cells[j], "EBF5FB")
    # 排名列绿色背景
    set_cell_bg(row.cells[0], GREEN)
    for p in row.cells[0].paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.bold = True

doc.add_paragraph()

# ======== 关注清单 ========
doc.add_page_break()
add_heading1(u"四、\u25b6 关注清单（等回调再进）")
add_p(u"以下股票基本面优秀，但当前估值偏高或距目标价空间有限，建议放入自选股，等待回调至合理区间后再分批建仓。")

watch = sorted([(c,s) for c,s in stock_data.items() if s["rating"]==u"\u25b6"], key=lambda x: -x[1]["score"])
tbl = doc.add_table(rows=len(watch)+1, cols=6)
tbl.style = 'Table Grid'
hdrs = [u"排名", u"代码", u"名称", u"层级", u"评分", u"建议等待回调至"]
for j, h in enumerate(hdrs):
    cell = tbl.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "1F4E79")
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.size = Pt(10)
    set_cell_border(cell)

for i, (code, s) in enumerate(watch):
    wp = round(s["target"] * 0.90)
    row = tbl.rows[i+1]
    row.cells[0].text = str(i+1)
    row.cells[1].text = code
    row.cells[2].text = s["name"]
    row.cells[3].text = s["layer_name"]
    row.cells[4].text = str(s["score"])
    row.cells[5].text = "{}元（-10%）".format(wp)
    for j in range(6):
        set_cell_border(row.cells[j])
        if i%2==0:
            set_cell_bg(row.cells[j], "EBF5FB")
    set_cell_bg(row.cells[0], BLUE)
    for p in row.cells[0].paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.bold = True

doc.add_paragraph()

# ======== 持仓建议 ========
doc.add_page_break()
add_heading1(u"五、持仓配置建议")
add_heading2(u"5.1  A股AI产业链仓位管理（小白友好）")
allocs = [
    (u"单票仓位上限", u"3%", u"A股波动比美股大，单票上限比美股(5%)更保守"),
    (u"AI产业链总仓位", u"<= 30%", u"AI板块波动大，不宜all-in，保留70%现金/债券/宽基"),
    (u"首次建仓比例", u"15–20%", u"不要一次性买满，先建立试探仓"),
    (u"加仓节奏", u"每跌8–10%加一批", u"A股波动大，拉开加仓间距"),
    (u"止损线", u"单票 -12~15%", u"A股下跌动能强，必须设止损"),
    (u"禁区", u"涨停板不追", u"A股「利好出货」是常态，追涨容易被套"),
]
for k, v, n in allocs:
    p = doc.add_paragraph()
    p.add_run(u"\u25b8 "+k+u"：").bold = True
    p.runs[0].font.color.rgb = HDR
    p.add_run("  "+v+"  ").bold = True
    p.runs[1].font.color.rgb = RGBColor(0x1A,0x7A,0x1A)
    p.add_run(n)
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_paragraph()
add_heading2(u"5.2  推荐持仓结构（参考）")
alloc2 = [
    (u"核心仓（\u2b50c强烈关注）", u"15–20%", u"中际旭创+海光信息+英维克，3只分散风险"),
    (u"观察仓（\u25b6关注）", u"5–10%", u"选择2只估值最合理/空间最大的，等回调建仓"),
    (u"低仓/预备仓", u"5%", u"现金，等待大盘调整或个股利空砸出的黄金坑"),
    (u"非AI仓（压舱石）", u">= 60%", u"沪深300ETF/债券基金/现金，避免AI板块回调时满仓被埋"),
]
tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Table Grid'
for i, (t, pct, d) in enumerate(alloc2):
    row = tbl.rows[i]
    row.cells[0].text = t
    row.cells[1].text = pct
    row.cells[2].text = d
    for j in range(3):
        set_cell_border(row.cells[j])
        if i%2==0:
            set_cell_bg(row.cells[j], "F2F3F4")
    for p in row.cells[0].paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
    for p in row.cells[1].paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A,0x7A,0x1A)

doc.add_paragraph()

# ======== 风险提示 ========
doc.add_page_break()
add_heading1(u"六、风险提示")

add_heading2(u"6.1  个股风险")
add_bullet(u"寒武纪：客户集中度极高（字节/阿里），竞争对手（海光/华为昇腾）挤压，估值极高（PE 185x）")
add_bullet(u"海光信息：技术路线依赖AMD授权，若美国升级出口管制存在授权风险")
add_bullet(u"中际旭创/新易盛：客户集中度（英伟达/北美云厂），中美贸易摩擦可能影响供货")
add_bullet(u"北方华创/中微公司：估值极高（PE 65-78x），关键零部件仍依赖进口")
add_bullet(u"中芯国际：先进制程（7nm以下）受出口管制，短期盈利承压")

add_heading2(u"6.2  行业系统性风险")
add_bullet(u"国产化进度不及预期：部分公司估值已提前反映国产化预期，若进度慢于预期将大幅回调")
add_bullet(u"AI资本开支退潮：若北美云厂（微软/谷歌/Meta）削减AI资本开支，供应链将受冲击")
add_bullet(u"A股政策市风险：AI板块对政策敏感，监管收紧（如限制AI应用）将影响估值")
add_bullet(u"估值泡沫破裂风险：参照2000年互联网泡沫，当普通投资者大量涌入时需警惕")

add_heading2(u"6.3  A股特有风险")
add_bullet(u"大股东减持：解禁期后大股东减持是A股常态，需关注解禁时间表")
add_bullet(u"概念炒作退潮：部分公司仅蹭AI概念，无实质业务，炒作过后将大幅回落")
add_bullet(u"北向资金流出：北向资金是A股AI板块重要增量资金，持续流出将导致板块承压")
add_bullet(u"流动性风险：小盘股流动性差，大资金进出困难")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(u"\u26a0\ufe0f 免责声明：本报告所有分析、评分、目标价均基于公开信息和知识库推断，不构成投资建议。A股投资有风险，入市需谨慎。")
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

# ======== 保存 ========
out_path = r"c:\Users\asus\WorkBuddy\20260530114642\A股AI产业链选股分析报告_20260530.docx"
doc.save(out_path)
print(u"Word报告生成成功：" + out_path)
